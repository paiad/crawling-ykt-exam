# @Software  :PyCharm
# @Author    : ad
# @File      : file_transform.py
# @Time      : 2025-03-01 13:24
import pandas as pd
from docx import Document

def txt_to_word(input_txt_path, output_docx_path):
    # 创建一个新的 Word 文档
    doc = Document()

    # 读取 .txt 文件内容
    with open(input_txt_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    # 将每一行写入 Word 文档
    for line in lines:
        # 去除每行末尾的换行符（因为 add_paragraph 会自动添加换行）
        line = line.strip()
        doc.add_paragraph(line)

    # 保存 Word 文档
    doc.save(output_docx_path)
    print(f"文件已成功转换为 {output_docx_path}")

exam_id = pd.read_csv('exam_id.csv')['exam_id'].iloc[0]
# 示例调用
input_txt_path = "./txt/雨课堂测试-id-{}.txt".format(exam_id) # 输入的 .txt 文件路径
output_docx_path = './docs/雨课堂测试-id-{}.docx'.format(exam_id)  # 输出的 .docx 文件路径
txt_to_word(input_txt_path, output_docx_path)