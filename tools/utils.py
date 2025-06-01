import re
from docx import Document
# 保存文件
def save_to_file(data, filename="data.txt"):
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(data + "\n\n")

# 使用正则表达式去除 HTML 标签
def remove_html_tags(text):
    clean_text = re.sub(r'<[^>]+>', '', text)
    return clean_text

# txt->docx
def txt_to_word(input_txt_path, output_docx_path):
    # 创建一个新的 Word 文档
    doc = Document()

    # 读取 .txt 文件内容
    with open(input_txt_path, 'r', encoding='utf-8') as file:
        lines = file.readlines()

    # 将每一行写入 Word 文档
    for line in lines:
        # 去除每行末尾的换行符（因为 add_paragraph 会自动添加换行）
        line = line.strip()
        doc.add_paragraph(line)

    # 保存 Word 文档
    doc.save(output_docx_path)
    print(f"文件已成功转换为 {output_docx_path}")
